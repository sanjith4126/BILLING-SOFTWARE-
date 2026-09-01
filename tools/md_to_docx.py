"""Convert USER_GUIDE.md to USER_GUIDE.docx.

Handles the subset of markdown we actually use in the guide:
- headings (#, ##, ###)
- horizontal rules (---)
- unordered lists (- item)
- ordered lists (1. item)
- pipe tables
- code fences (```)
- bold **text** and inline `code`

Deliberately simple: python-docx only, no pandoc, no external deps.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")


def add_runs(paragraph, text: str) -> None:
    """Add runs to paragraph, honouring **bold** and `code` inline markup."""
    # Tokenize: split on **...** and `...` while preserving order.
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        chunk = m.group(0)
        if chunk.startswith("**"):
            r = paragraph.add_run(chunk[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line))


def split_table_row(line: str) -> list[str]:
    parts = line.strip().split("|")
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return [p.strip() for p in parts]


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule → thin bordered paragraph
        if re.match(r"^-{3,}$", line.strip()):
            p = doc.add_paragraph()
            r = p.add_run("_" * 60)
            r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt({1: 20, 2: 16, 3: 13}.get(level, 12))
            if level == 1:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        # Code fence
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        # Tables (pipe format)
        if "|" in line and i + 1 < n and is_table_separator(lines[i + 1]):
            header = split_table_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, h in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                add_runs(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r_idx, row in enumerate(rows, start=1):
                for c, val in enumerate(row):
                    if c >= len(header):
                        continue
                    cell = table.rows[r_idx].cells[c]
                    cell.text = ""
                    add_runs(cell.paragraphs[0], val)
            doc.add_paragraph()
            continue

        # Ordered list item
        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # Unordered list item (- or *)
        m = re.match(r"^[-*]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue

        # Nested list item (2-space indent)
        m = re.match(r"^\s{2,}[-*]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet 2")
            add_runs(p, m.group(1))
            i += 1
            continue

        # Plain paragraph — merge consecutive non-blank lines that aren't special
        para_lines = [line.rstrip()]
        j = i + 1
        while j < n:
            nxt = lines[j]
            if (not nxt.strip()
                    or re.match(r"^#{1,6}\s", nxt)
                    or re.match(r"^-{3,}$", nxt.strip())
                    or re.match(r"^(\d+)\.\s", nxt)
                    or re.match(r"^[-*]\s", nxt)
                    or nxt.strip().startswith("```")
                    or ("|" in nxt and j + 1 < n and is_table_separator(lines[j + 1]))):
                break
            para_lines.append(nxt.rstrip())
            j += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(l.strip() for l in para_lines))
        i = j

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    md = root / "USER_GUIDE.md"
    docx = root / "USER_GUIDE.docx"
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
    if len(sys.argv) > 2:
        docx = Path(sys.argv[2])
    convert(md, docx)
